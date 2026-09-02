"""Export the Gate 17A Vietnamese report to DOCX and PDF.

The exporter consumes the already reviewed Markdown report. It does not run
simulation, calibration, tuning, or alter historical metrics/manifests.
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
import hashlib
import html
import json
from pathlib import Path
import platform
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs/report/final_vietnamese_report.md"
EXPORT_ROOT = ROOT / "experiments/gate_17b_report_export"
RESULTS_ROOT = EXPORT_ROOT / "results"
MANIFEST_ROOT = EXPORT_ROOT / "manifests"
DOCX_PATH = RESULTS_ROOT / "final_vietnamese_report.docx"
PDF_PATH = RESULTS_ROOT / "final_vietnamese_report.pdf"
SUMMARY_PATH = RESULTS_ROOT / "report_export_summary.json"
MANIFEST_PATH = MANIFEST_ROOT / "report_export_manifest.json"
EXPORT_README = ROOT / "docs/report/export_readme.md"

FINAL_CLAIM = (
    "Chen-calibrated organism-level computational locomotion proxy with "
    "directional Pozo holdout concordance and substantial quantitative ratio mismatch"
)


@dataclass(frozen=True)
class Block:
    kind: str
    text: str = ""
    level: int = 0
    rows: tuple[tuple[str, ...], ...] = ()


def _require(condition: bool, message: str) -> None:
    if not condition:
        raise ValueError(message)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _relative(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def _git_value(*args: str) -> str:
    result = subprocess.run(
        ["git", *args],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout.strip()


def _write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def _split_table_row(line: str) -> tuple[str, ...]:
    stripped = line.strip().strip("|")
    return tuple(part.strip() for part in stripped.split("|"))


def parse_markdown(path: Path) -> list[Block]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if paragraph:
            text = " ".join(part.strip() for part in paragraph).strip()
            if text:
                blocks.append(Block(kind="paragraph", text=text))
            paragraph.clear()

    while index < len(lines):
        line = lines[index]
        if line.strip().startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            _require(index < len(lines), "Unclosed Markdown code fence")
            blocks.append(Block(kind="code", text="\n".join(code_lines)))
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+?)\s*$", line)
        if heading:
            flush_paragraph()
            blocks.append(Block(kind="heading", level=len(heading.group(1)), text=heading.group(2)))
            index += 1
            continue

        if (
            line.strip().startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1])
        ):
            flush_paragraph()
            rows = [_split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            blocks.append(Block(kind="table", rows=tuple(rows)))
            continue

        bullet = re.match(r"^\s*-\s+(.+?)\s*$", line)
        if bullet:
            flush_paragraph()
            blocks.append(Block(kind="bullet", text=bullet.group(1)))
            index += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+?)\s*$", line)
        if numbered:
            flush_paragraph()
            blocks.append(Block(kind="number", text=numbered.group(1)))
            index += 1
            continue

        if not line.strip():
            flush_paragraph()
        else:
            paragraph.append(line)
        index += 1

    flush_paragraph()
    return blocks


def _set_font(run: Any, name: str, size: float, color: str = "000000", bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    normalized = color.removeprefix("#")
    run.font.color.rgb = RGBColor(
        int(normalized[0:2], 16),
        int(normalized[2:4], 16),
        int(normalized[4:6], 16),
    )
    run.bold = bold


def _add_inline_docx(paragraph: Any, text: str, size: float = 11, color: str = "000000") -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            _set_font(run, "Calibri", size, color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_font(run, "Calibri", size, color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            _set_font(run, "Consolas", size - 0.3, color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_font(run, "Calibri", size, color)


def _set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell: Any, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        normalized = color.removeprefix("#")
        style.font.color.rgb = RGBColor(
            int(normalized[0:2], 16),
            int(normalized[2:4], 16),
            int(normalized[4:6], 16),
        )
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = document.styles.add_style("Report Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Gate 17B | Báo cáo khoa học tiếng Việt")
    _set_font(header_run, "Calibri", 9, "6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Trang ")
    _set_font(footer_run, "Calibri", 9, "6B7280")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    footer_run._r.append(begin)
    footer_run._r.append(instruction)
    footer_run._r.append(end)


def _add_docx_table(document: Document, rows: tuple[tuple[str, ...], ...]) -> None:
    column_count = max(len(row) for row in rows)
    normalized = [row + ("",) * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    if column_count == 2:
        widths = [2700, 6660]
    else:
        widths = [9360 // column_count] * column_count
        widths[-1] += 9360 - sum(widths)
    _set_table_geometry(table, widths)
    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            _add_inline_docx(paragraph, value, size=9.5, color="000000")
            if row_index == 0:
                _shade_cell(cell, "F2F4F7")
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx(blocks: Iterable[Block], output: Path) -> None:
    document = Document()
    document.core_properties.title = "Báo cáo tổng kết dự án Drosophila Parkinson-like Locomotion Proxy"
    document.core_properties.subject = "Gate 17B report export"
    document.core_properties.language = "vi-VN"
    _configure_docx(document)

    first_heading = True
    for block in blocks:
        if block.kind == "heading":
            if first_heading and block.level == 1:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(8)
                run = paragraph.add_run(block.text)
                _set_font(run, "Calibri", 24, "0B2545", bold=True)
                first_heading = False
                subtitle = document.add_paragraph()
                subtitle.paragraph_format.space_after = Pt(18)
                subtitle_run = subtitle.add_run(
                    "Bản xuất Gate 17B từ docs/report/final_vietnamese_report.md | release v1.0.0"
                )
                _set_font(subtitle_run, "Calibri", 10, "6B7280")
            else:
                paragraph = document.add_paragraph(style=f"Heading {block.level}")
                _add_inline_docx(paragraph, block.text, size={1: 16, 2: 13, 3: 12}[block.level], color="2E74B5")
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_inline_docx(paragraph, block.text)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_docx(paragraph, block.text)
        elif block.kind == "number":
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_docx(paragraph, block.text)
        elif block.kind == "code":
            for line in block.text.splitlines() or [""]:
                paragraph = document.add_paragraph(style="Report Code")
                run = paragraph.add_run(line)
                _set_font(run, "Consolas", 8.5, "374151")
        elif block.kind == "table":
            _add_docx_table(document, block.rows)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def _register_pdf_fonts() -> tuple[str, str]:
    candidates = [
        (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")),
        (Path("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"), Path("/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf")),
    ]
    for regular, bold in candidates:
        if regular.is_file() and bold.is_file():
            pdfmetrics.registerFont(TTFont("ReportArial", str(regular)))
            pdfmetrics.registerFont(TTFont("ReportArial-Bold", str(bold)))
            return "ReportArial", "ReportArial-Bold"
    raise RuntimeError("Không tìm thấy font TrueType hỗ trợ tiếng Việt cho PDF")


def _inline_pdf(text: str) -> str:
    escaped = html.escape(text, quote=False)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"`(.+?)`", r'<font name="ReportMono">\1</font>', escaped)
    return escaped


def _hex_color(value: str) -> Any:
    return colors.HexColor("#" + value.removeprefix("#"))


def _pdf_styles(font: str, bold_font: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()["BodyText"]
    return {
        "title": ParagraphStyle("ReportTitle", parent=base, fontName=bold_font, fontSize=24, leading=29, textColor=_hex_color("0B2545"), spaceAfter=8),
        "subtitle": ParagraphStyle("ReportSubtitle", parent=base, fontName=font, fontSize=9.5, leading=13, textColor=_hex_color("6B7280"), spaceAfter=18),
        "h1": ParagraphStyle("ReportH1", parent=base, fontName=bold_font, fontSize=16, leading=20, textColor=_hex_color("2E74B5"), spaceBefore=18, spaceAfter=10, keepWithNext=True),
        "h2": ParagraphStyle("ReportH2", parent=base, fontName=bold_font, fontSize=13, leading=17, textColor=_hex_color("2E74B5"), spaceBefore=12, spaceAfter=6, keepWithNext=True),
        "h3": ParagraphStyle("ReportH3", parent=base, fontName=bold_font, fontSize=12, leading=15, textColor=_hex_color("1F4D78"), spaceBefore=8, spaceAfter=4, keepWithNext=True),
        "body": ParagraphStyle("ReportBody", parent=base, fontName=font, fontSize=10.5, leading=14, alignment=TA_JUSTIFY, spaceAfter=8),
        "bullet": ParagraphStyle("ReportBullet", parent=base, fontName=font, fontSize=10.5, leading=14, leftIndent=16, firstLineIndent=-9, spaceAfter=4),
        "number": ParagraphStyle("ReportNumber", parent=base, fontName=font, fontSize=10.5, leading=14, leftIndent=18, firstLineIndent=-12, spaceAfter=4),
        "table": ParagraphStyle("ReportTable", parent=base, fontName=font, fontSize=8.7, leading=11, alignment=TA_LEFT),
        "table_header": ParagraphStyle("ReportTableHeader", parent=base, fontName=bold_font, fontSize=8.7, leading=11, alignment=TA_LEFT),
        "code": ParagraphStyle("ReportCode", parent=base, fontName="ReportMono", fontSize=7.8, leading=10, textColor=_hex_color("374151"), leftIndent=8, rightIndent=8, spaceAfter=0),
    }


def _pdf_page_number(canvas: Any, document: Any) -> None:
    canvas.saveState()
    canvas.setFont("ReportArial", 8.5)
    canvas.setFillColor(_hex_color("6B7280"))
    canvas.drawRightString(7.5 * inch, 0.55 * inch, f"Trang {document.page}")
    canvas.drawRightString(7.5 * inch, 10.45 * inch, "Gate 17B | Báo cáo khoa học tiếng Việt")
    canvas.restoreState()


def build_pdf(blocks: Iterable[Block], output: Path) -> int:
    font, bold_font = _register_pdf_fonts()
    pdfmetrics.registerFont(TTFont("ReportMono", "C:/Windows/Fonts/consola.ttf")) if Path("C:/Windows/Fonts/consola.ttf").is_file() else None
    styles = _pdf_styles(font, bold_font)
    document = SimpleDocTemplate(
        str(output),
        pagesize=LETTER,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Báo cáo tổng kết dự án",
        author="Drosophila Parkinson Neural Disease project",
    )
    story: list[Any] = []
    first_heading = True
    for block in blocks:
        if block.kind == "heading":
            if first_heading and block.level == 1:
                story.append(Paragraph(_inline_pdf(block.text), styles["title"]))
                story.append(Paragraph("Bản xuất Gate 17B từ docs/report/final_vietnamese_report.md | release v1.0.0", styles["subtitle"]))
                first_heading = False
            else:
                story.append(Paragraph(_inline_pdf(block.text), styles[f"h{block.level}"]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_inline_pdf(block.text), styles["body"]))
        elif block.kind == "bullet":
            story.append(Paragraph("&#8226;&nbsp;" + _inline_pdf(block.text), styles["bullet"]))
        elif block.kind == "number":
            story.append(Paragraph("&#8226;&nbsp;" + _inline_pdf(block.text), styles["number"]))
        elif block.kind == "code":
            story.append(KeepTogether([Preformatted(block.text, styles["code"]), Spacer(1, 6)]))
        elif block.kind == "table":
            data = []
            max_columns = max(len(row) for row in block.rows)
            for row_index, row in enumerate(block.rows):
                normalized = row + ("",) * (max_columns - len(row))
                style = styles["table_header"] if row_index == 0 else styles["table"]
                data.append([Paragraph(_inline_pdf(value), style) for value in normalized])
            available = 6.9 * inch
            col_widths = [2.0 * inch, 4.9 * inch] if max_columns == 2 else [available / max_columns] * max_columns
            table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _hex_color("F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, _hex_color("B8C1CC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.extend([Spacer(1, 3), table, Spacer(1, 8)])
    output.parent.mkdir(parents=True, exist_ok=True)
    document.build(story, onFirstPage=_pdf_page_number, onLaterPages=_pdf_page_number)
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz

        pdf = fitz.open(str(output))
        page_count = len(pdf)
        qa_dir = ROOT / "temporary/gate_17b_report_export_visual_qa"
        qa_dir.mkdir(parents=True, exist_ok=True)
        for page_index, page in enumerate(pdf):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            pixmap.save(str(qa_dir / f"page-{page_index + 1:03d}.png"))
        pdf.close()
        return page_count
    except ImportError:
        return -1


def _visual_qa_status() -> tuple[str, str]:
    has_docx_renderer = shutil.which("soffice") is not None or shutil.which("libreoffice") is not None
    qa_dir = ROOT / "temporary/gate_17b_report_export_visual_qa"
    has_pdf_renders = qa_dir.is_dir() and any(qa_dir.glob("page-*.png"))
    if has_docx_renderer and has_pdf_renders:
        return "PASS", "DOCX renderer and PDF page renders are available."
    if not has_docx_renderer:
        return "PARTIAL_DOCX_RENDERER_UNAVAILABLE", "LibreOffice/soffice is not installed; DOCX visual render was not available."
    return "PARTIAL_PDF_RENDER_UNAVAILABLE", "PDF page render artifacts were not available."


def export_package(source: Path) -> dict[str, Any]:
    _require(source.is_file(), f"Missing report source: {_relative(source)}")
    _require(EXPORT_README.is_file(), f"Missing export README: {_relative(EXPORT_README)}")
    blocks = parse_markdown(source)
    _require(any(block.kind == "heading" for block in blocks), "Report has no headings")
    _require(any(block.kind == "table" for block in blocks), "Report has no tables")
    build_docx(blocks, DOCX_PATH)
    page_count = build_pdf(blocks, PDF_PATH)
    visual_status, visual_note = _visual_qa_status()
    status = "REPORT_EXPORT_READY" if visual_status == "PASS" else "REPORT_EXPORT_PARTIAL"
    summary = {
        "schema_version": "gate-17b-report-export-summary-v1",
        "status": status,
        "source_markdown": _relative(source),
        "docx": _relative(DOCX_PATH),
        "pdf": _relative(PDF_PATH),
        "pdf_page_count": page_count,
        "visual_qa_status": visual_status,
        "visual_qa_note": visual_note,
        "claim_lock_active": True,
        "final_claim": FINAL_CLAIM,
        "no_new_simulation_run": True,
        "no_calibration_run": True,
        "no_tuning_run": True,
        "no_raw_metric_modification": True,
        "boundaries": {
            "biological_validation": False,
            "gene_specific_validation": False,
            "clinical_validation": False,
            "drug_validation": False,
            "therapeutic_validation": False,
        },
    }
    _write_json(SUMMARY_PATH, summary)
    hash_files = [source, EXPORT_README, DOCX_PATH, PDF_PATH, SUMMARY_PATH]
    manifest = {
        "schema_version": "gate-17b-report-export-manifest-v1",
        "status": status,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "git_commit": _git_value("rev-parse", "HEAD"),
        "python_version": platform.python_version(),
        "source_files": [_relative(source), _relative(EXPORT_README)],
        "generated_files": [
            _relative(DOCX_PATH),
            _relative(PDF_PATH),
            _relative(SUMMARY_PATH),
        ],
        "sha256": {_relative(path): _sha256(path) for path in hash_files},
        "visual_qa_status": visual_status,
        "visual_qa_note": visual_note,
        "claim_lock_active": True,
        "no_new_simulation_run": True,
        "no_calibration_run": True,
        "no_tuning_run": True,
        "no_raw_metric_modification": True,
        "large_artifacts_committed": False,
    }
    _write_json(MANIFEST_PATH, manifest)
    return summary


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    args = parser.parse_args(argv)
    try:
        summary = export_package(args.source)
    except (OSError, RuntimeError, subprocess.SubprocessError, ValueError, KeyError, TypeError) as exc:
        print(f"Report export failed: {exc}", file=sys.stderr)
        return 1
    print(f"Status: {summary['status']}")
    print(f"DOCX: {DOCX_PATH}")
    print(f"PDF: {PDF_PATH}")
    print(f"Summary: {SUMMARY_PATH}")
    print(f"Manifest: {MANIFEST_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
